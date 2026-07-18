import streamlit as st
import google.generativeai as genai
import os
from pathlib import Path
import json
import re
import io
import time
import random
import base64
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import matplotlib
matplotlib.use('Agg')  # Non-interactive backend for server
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import matplotlib.ticker as ticker
import numpy as np
from fpdf import FPDF

# Helper: Clean text for export (removes Hangul and Emojis)
def clean_text_for_export(text):
    if not text: return ""
    # Export fonts (standard PDF/Word) handle Latin-1 best. 
    # Encoding to latin-1 with 'ignore' strips Hangul and Emojis.
    return text.encode('latin-1', 'ignore').decode('latin-1')

# Load environment variables
load_dotenv()

# Page Configuration
st.set_page_config(
    page_title="Elite Prep | SAT Analysis & Adaptive Practice",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Brand Colors & Styles
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');
    
    html, body, [class*="css"]  {
        font-family: 'Inter', sans-serif;
        color: #1f2937;
        background-color: #f3f4f6;
    }
    
    /* Main Container */
    .block-container {
        padding-top: 2rem;
        padding-bottom: 5rem;
        max-width: 1000px;
    }

    /* Header Styling */
    .main-header {
        font-family: 'Inter', sans-serif;
        color: #0C1E41; /* Elite Navy */
        text-align: center;
        font-weight: 800;
        font-size: 2.2rem;
        margin-bottom: 0.5rem;
        letter-spacing: -0.02em;
    }
    .sub-header {
        color: #6b7280;
        text-align: center;
        margin-bottom: 3rem;
        font-size: 1.1rem;
        font-weight: 400;
    }
    
    /* Card Styling */
    .stCard {
        background-color: white;
        padding: 2rem;
        border-radius: 12px;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
        margin-bottom: 2rem;
        border: 1px solid #e5e7eb;
    }

    /* Button Styling */
    div.stButton > button {
        background-color: #0C1E41;
        color: white;
        font-weight: 600;
        border-radius: 8px;
        padding: 0.75rem 1.5rem;
        border: none;
        transition: all 0.2s ease;
        width: 100%;
        box-shadow: 0 4px 6px rgba(12, 30, 65, 0.2);
    }
    div.stButton > button:hover {
        background-color: #1a3a6e;
        transform: translateY(-1px);
        box-shadow: 0 6px 8px rgba(12, 30, 65, 0.3);
    }

    /* Practice Topic Buttons (Secondary) */
    .topic-btn {
        border: 1px solid #0C1E41 !important;
        background-color: white !important;
        color: #0C1E41 !important;
    }
    
    /* Status Indicators */
    .status-success { color: #166534; font-weight: 500;}
    .status-error { color: #991b1b; font-weight: 500;}

</style>
""", unsafe_allow_html=True)

# Helper: Upload file to Gemini (File API)
def upload_to_gemini(path, mime_type=None):
    try:
        file = genai.upload_file(path, mime_type=mime_type)
        return file
    except Exception as e:
        print(f"Error uploading {path}: {e}")
        return None

# Candidate models tried in order. "-latest" aliases always point to a live
# model, so they keep working even when a pinned version (e.g. gemini-2.5-pro)
# is retired for a given API key. The pinned names are kept as extra fallbacks
# for keys/projects that only expose those.
GEMINI_MODEL_CANDIDATES = [
    'gemini-flash-latest',
    'gemini-pro-latest',
    'gemini-2.5-flash',
    'gemini-2.5-pro',
]


def _is_model_unavailable_error(msg):
    """True when the error means 'this model name can't be used with this key'."""
    m = msg.lower()
    return (
        '404' in msg
        or 'not found' in m
        or 'no longer available' in m
        or 'is not available' in m
        or 'not supported' in m
        or 'does not exist' in m
    )


# Helper: Load Gemini Model
def get_gemini_response(input_prompt, content_parts, temperature=0.2):
    try:
        api_key = os.getenv("GOOGLE_API_KEY")
        if not api_key:
            if 'api_key' in st.session_state and st.session_state.api_key:
                api_key = st.session_state.api_key
            else:
                st.error("Google API Key not found.")
                return None

        genai.configure(api_key=api_key)

        full_payload = [input_prompt] + content_parts

        generation_config = genai.types.GenerationConfig(
            temperature=temperature,
            max_output_tokens=32768,
        )

        # Try a model that already worked this session first, then the rest.
        candidates = list(GEMINI_MODEL_CANDIDATES)
        remembered = st.session_state.get('working_model')
        if remembered and remembered in candidates:
            candidates.remove(remembered)
            candidates.insert(0, remembered)

        unavailable = []
        for model_name in candidates:
            model = genai.GenerativeModel(model_name)

            # Retry on transient errors (e.g. 504 timeout / 503 overloaded).
            # These are usually temporary, so we automatically retry a few times
            # instead of forcing the user to press the button again.
            max_retries = 3
            skip_to_next_model = False
            for attempt in range(max_retries):
                try:
                    response = model.generate_content(
                        full_payload,
                        generation_config=generation_config,
                        request_options={"timeout": 600},
                    )

                    # Safe access to text
                    try:
                        text = response.text
                        st.session_state['working_model'] = model_name
                        return text
                    except ValueError:
                        # Handle cases where response is blocked or empty
                        finish_reason = response.candidates[0].finish_reason if response.candidates else "UNKNOWN"
                        st.error(f"Generation stopped. Finish Reason: {finish_reason}")
                        # If MAX_TOKENS (2), try to return what we have
                        if finish_reason == 2 and response.candidates and response.candidates[0].content.parts:
                            st.session_state['working_model'] = model_name
                            return response.candidates[0].content.parts[0].text
                        return None
                except Exception as e:
                    msg = str(e)
                    # Model not usable with this key -> fall back to the next one.
                    if _is_model_unavailable_error(msg):
                        unavailable.append(model_name)
                        skip_to_next_model = True
                        break
                    # Only retry transient/timeout errors; fail fast on others.
                    is_transient = any(code in msg for code in ("504", "503", "500", "timed out", "timeout", "deadline"))
                    if is_transient and attempt < max_retries - 1:
                        wait = 3 * (attempt + 1)
                        st.warning(f"⏳ Request timed out (attempt {attempt + 1}/{max_retries}). Retrying in {wait}s...")
                        time.sleep(wait)
                        continue
                    # Non-transient, non-availability error -> genuine failure.
                    raise

            if skip_to_next_model:
                continue

        # Every candidate model was unavailable for this key.
        raise RuntimeError(
            "None of the configured Gemini models are available for this API key "
            f"(tried: {', '.join(candidates)}). The models may have been retired — "
            "update GEMINI_MODEL_CANDIDATES in app.py."
        )
    except Exception as e:
        st.error(f"Error calling Gemini API: {str(e)}")
        return None

# Helper: Extract JSON from text
def extract_json_topics(text):
    try:
        # Find JSON block using regex
        match = re.search(r'```json\s*(\{.*?\})\s*```', text, re.DOTALL)
        if match:
            return json.loads(match.group(1))
        else:
            # Try finding just brace content if no markdown tags
            match = re.search(r'\{.*\}', text, re.DOTALL)
            if match:
                return json.loads(match.group(0))
    except Exception as e:
        print(f"JSON Parse Error: {e}")
    return None

# Helper: Add colored difficulty badges to rendered output
def add_difficulty_badges(text):
    """Replace [Easy], [Medium], [Hard] tags with colored HTML badge spans."""
    if not text:
        return text
    text = text.replace('[Easy]',   '<span style="background:#22c55e;color:white;padding:2px 9px;border-radius:12px;font-size:0.78em;font-weight:700;margin-right:5px">Easy</span>')
    text = text.replace('[Medium]', '<span style="background:#f59e0b;color:white;padding:2px 9px;border-radius:12px;font-size:0.78em;font-weight:700;margin-right:5px">Medium</span>')
    text = text.replace('[Hard]',   '<span style="background:#ef4444;color:white;padding:2px 9px;border-radius:12px;font-size:0.78em;font-weight:700;margin-right:5px">Hard</span>')
    return text

# Helper: Decide if a $...$ span is really math or just prose the model
# wrongly wrapped in math delimiters.
def _is_prose_span(content):
    # Remove \text{...} blocks and \commands (frac, cap, sqrt, ...) — these are
    # legitimate math and should not count as English words.
    stripped = re.sub(r'\\text\{[^}]*\}', ' ', content)
    stripped = re.sub(r'\\[a-zA-Z]+', ' ', stripped)
    # Real inline math almost never has 2+ multi-letter English words left.
    words = re.findall(r'[A-Za-z]{2,}', stripped)
    return len(words) >= 2

# Helper: Safety net for malformed LaTeX coming from the model.
# The model sometimes (a) leaves an unbalanced $ or (b) wraps whole English
# sentences in $...$. Both break Streamlit's KaTeX rendering (text runs
# together in math italics). This normalizes the text so it always renders
# readably, even if a formula occasionally shows as raw LaTeX.
def sanitize_math_delimiters(text):
    if not text:
        return text
    out_lines = []
    for line in text.split('\n'):
        # Leave display math ($$...$$) and lines without $ untouched.
        if '$$' in line:
            out_lines.append(line)
            continue
        positions = [m.start() for m in re.finditer(r'(?<!\\)\$', line)]
        if not positions:
            out_lines.append(line)
            continue
        if len(positions) % 2 != 0:
            # Unbalanced $ on this line -> neutralize all so it renders as text.
            out_lines.append(line.replace('$', r'\$'))
            continue
        # Balanced: unwrap any pair whose content is actually prose.
        remove_idx = set()
        for k in range(0, len(positions), 2):
            start, end = positions[k], positions[k + 1]
            if _is_prose_span(line[start + 1:end]):
                remove_idx.add(start)
                remove_idx.add(end)
        if remove_idx:
            line = ''.join(ch for i, ch in enumerate(line) if i not in remove_idx)
        out_lines.append(line)
    return '\n'.join(out_lines)

# Helper: Force Formatting for Options
def ensure_formatting(text):
    if not text: return ""
    # Remove <br> tags
    text = re.sub(r'<br\s*/?>', '', text)
    # Regex to find A) B) C) D) that are NOT at start of line
    pattern = r"(\s)([A-D]\)) "
    formatted_text = re.sub(pattern, r"\n\n\2 ", text)
    return formatted_text

# Helper: Execute matplotlib code and return PNG image bytes
def execute_figure_code(code_str):
    """Execute matplotlib code string and return PNG image as bytes."""
    try:
        plt.close('all')  # Close any previous figures
        
        # Pre-process the code to fix common issues
        lines = code_str.split('\n')
        cleaned_lines = []
        for line in lines:
            stripped = line.strip()
            # Skip import statements (plt, np, io, math already available)
            if stripped.startswith('import ') or stripped.startswith('from '):
                continue
            # Replace plt.show() with savefig
            if stripped == 'plt.show()':
                continue
            # Drop the model's own savefig — we save ourselves after fixing the
            # legend, so the saved image never has an overlapping legend.
            if 'savefig' in stripped:
                continue
            cleaned_lines.append(line)

        cleaned_code = '\n'.join(cleaned_lines)

        # Create a BytesIO buffer for the image
        buf = io.BytesIO()
        
        # Set up execution namespace with common imports available
        exec_namespace = {
            'plt': plt,
            'np': np,
            'io': io,
            'buf': buf,
            'math': __import__('math'),
            'matplotlib': matplotlib,
            'patches': patches,
            'ticker': ticker,
        }
        
        # Execute the code
        exec(cleaned_code, exec_namespace)

        # Move any legend outside the plot area so it never overlaps the data,
        # then save the figure ourselves (the model's savefig was stripped).
        fig = plt.gcf()
        if fig.get_axes():
            for ax in fig.get_axes():
                if ax.get_legend() is not None:
                    handles, labels = ax.get_legend_handles_labels()
                    if handles:
                        ax.legend(
                            handles, labels,
                            loc='upper left', bbox_to_anchor=(1.02, 1.0),
                            borderaxespad=0.0, fontsize=9, framealpha=0.95,
                        )
            fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                        facecolor='white', edgecolor='none')

        buf.seek(0)
        img_bytes = buf.getvalue()
        plt.close('all')
        
        if len(img_bytes) > 0:
            return img_bytes
        return None
    except Exception as e:
        plt.close('all')
        # Silence warning as requested by user
        # st.warning(f"⚠️ Figure generation error: {e}")
        return None

# Helper: Parse Gemini response into text and figure segments
def parse_response_with_figures(response_text):
    """
    Parse response containing ```python-figure code blocks.
    Returns list of segments: [{"type":"text","content":"..."}, {"type":"figure","image":bytes,"code":"..."}]
    """
    if not response_text:
        return [{"type": "text", "content": ""}]
    
    # Pattern to match ```python-figure ... ``` OR ```python ... ``` blocks containing plt
    pattern = r'```(?:python-figure|python)\s*\n(.*?)```'
    
    segments = []
    last_end = 0
    
    for match in re.finditer(pattern, response_text, re.DOTALL):
        # Add text before this code block
        text_before = response_text[last_end:match.start()].strip()
        if text_before:
            segments.append({"type": "text", "content": text_before})
        
        # Execute the figure code
        code = match.group(1).strip()
        img_bytes = execute_figure_code(code)
        
        if img_bytes:
            segments.append({"type": "figure", "image": img_bytes, "code": code})
        else:
            # If figure generation failed, add a note
            segments.append({"type": "text", "content": "*[Figure could not be generated]*"})
        
        last_end = match.end()
    
    # Add remaining text after last code block
    remaining = response_text[last_end:].strip()
    if remaining:
        segments.append({"type": "text", "content": remaining})
    
    # If no figures found, return entire text as single segment
    if not segments:
        segments.append({"type": "text", "content": response_text})
    
    return segments

# Helper: Convert LaTeX math to readable plain text
def _latex_to_readable(text):
    """Convert LaTeX math expressions to readable Unicode text."""
    if not text:
        return text
    
    def process_math(match):
        math = match.group(1)
        # Handle nested fractions (inner to outer)
        max_iter = 10
        while r'\frac' in math and max_iter > 0:
            math = re.sub(r'\\frac\{([^{}]*)\}\{([^{}]*)\}', r'(\1)/(\2)', math)
            max_iter -= 1
        # Common LaTeX symbols
        replacements = {
            r'\times': '×', r'\div': '÷', r'\pm': '±',
            r'\leq': '≤', r'\geq': '≥', r'\neq': '≠',
            r'\rightarrow': '→', r'\leftarrow': '←',
            # Drop \left / \right sizing commands (must come AFTER the arrow
            # replacements above, since \left/\right are prefixes of
            # \leftarrow/\rightarrow, and BEFORE the \le/\ge rules below).
            r'\left': '', r'\right': '',
            r'\cdot': '·', r'\approx': '≈', r'\infty': '∞',
            r'\pi': 'π', r'\theta': 'θ', r'\alpha': 'α', r'\beta': 'β',
            r'\sqrt': '√', r'\le': '≤', r'\ge': '≥',
        }
        for latex_cmd, symbol in replacements.items():
            math = math.replace(latex_cmd, symbol)
        # Subscripts: Try common ones first
        subscripts = {'0': '₀', '1': '₁', '2': '₂', '3': '₃', '4': '₄', '5': '₅', '6': '₆', '7': '₇', '8': '₈', '9': '₉'}
        for k, v in subscripts.items():
            math = math.replace('_{' + k + '}', v).replace('_' + k, v)
        
        # Superscripts: Handle common ones (Latin-1 compatible for PDF)
        superscripts = {'1': '¹', '2': '²', '3': '³'}
        for k, v in superscripts.items():
            math = math.replace('^{' + k + '}', v).replace('^' + k, v)
            
        # Fallback for other superscripts/subscripts
        math = re.sub(r'_\{([^}]*)\}', r'_\1', math)
        math = re.sub(r'\^\{([^}]*)\}', r'^(\1)', math)
        
        # Remove remaining backslashes from unknown commands
        math = re.sub(r'\\([a-zA-Z]+)\s*', r'\1 ', math)
        # Clean extra spaces
        math = re.sub(r'\s+', ' ', math).strip()
        return math
    
    # Process $$...$$ blocks first, then $...$
    text = re.sub(r'\$\$\s*([^$]+?)\s*\$\$', process_math, text)
    text = re.sub(r'\$([^$]+?)\$', process_math, text)
    
    # Also convert bare LaTeX commands outside $ delimiters
    while r'\frac' in text:
        prev = text
        text = re.sub(r'\\frac\{([^{}]*)\}\{([^{}]*)\}', r'(\1)/(\2)', text)
        if text == prev:
            break
    
    bare_replacements = {
        r'\times': '×', r'\cdot': '·', r'\rightarrow': '→',
        r'\left': '', r'\right': '',
        r'\approx': '≈', r'\leq': '≤', r'\geq': '≥', r'\neq': '≠',
        r'\pm': '±', r'\div': '÷',
    }
    for latex_cmd, symbol in bare_replacements.items():
        text = text.replace(latex_cmd, symbol)
    
    # Convert bare superscripts for export
    text = text.replace('^{2}', '²').replace('^2', '²')
    text = text.replace('^{3}', '³').replace('^3', '³')
    text = text.replace('^{1}', '¹').replace('^1', '¹')
    
    return text

# ---------------------------------------------------------------------------
# Elite Prep styled Word export (matches sample-word-file.docx)
# ---------------------------------------------------------------------------

# Brand palette (hex, no leading #)
_ELITE_NAVY   = '14315B'   # navy — used for headings/text (not as a fill:
                           # solid dark banners burn printer toner)
_ELITE_MUTED  = '5A6B80'   # muted gray-blue for subtitles
_ELITE_LIGHT  = 'F2F5FA'   # light panel background
_ELITE_BORDER = 'D7E0EC'   # subtle card border
_ELITE_OPTION = '1E4E86'   # option letter blue
_ELITE_ANSWER = '2E6DB4'   # answer letter blue
# Difficulty badges: pale tint + dark text of the same hue. The colour still
# separates Easy/Medium/Hard at a glance, but a saturated fill with white text
# would lay down far more toner.
_BADGE = {
    'EASY':   {'fill': 'DCE9F7', 'text': '1B4F8A'},
    'MEDIUM': {'fill': 'FFF1C9', 'text': '7A5200'},
    'HARD':   {'fill': 'FBDDD9', 'text': 'A32A1C'},
}
_CARD_WIDTH  = 10360   # twips (~7.19")
_BADGE_COL   = 1100    # twips (~0.76")
_BODY_COL    = 9260    # twips (~6.43")


def _shade_cell(cell, fill_hex):
    """Apply a solid background fill to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _row_cant_split(row):
    """Prevent a table row from breaking across a page boundary."""
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn('w:cantSplit')) is None:
        trPr.append(OxmlElement('w:cantSplit'))


def _keep_with_next(paragraph):
    """Keep this paragraph on the same page as the following content."""
    paragraph.paragraph_format.keep_with_next = True


def _table_borders(table, color=_ELITE_BORDER, show=True):
    """Set uniform borders (or remove them) on a table."""
    tblPr = table._tbl.tblPr
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        if show:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), color)
        else:
            el.set(qn('w:val'), 'none')
            el.set(qn('w:sz'), '0')
            el.set(qn('w:color'), 'auto')
        borders.append(el)
    tblPr.append(borders)


def _table_bottom_rule(table, color=_ELITE_NAVY, sz='12'):
    """
    Draw only a bottom rule under a table (no fill, no other edges).
    Used instead of a solid dark banner so headings stay legible without
    covering the page in toner-hungry ink.
    """
    tblPr = table._tbl.tblPr
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        if edge == 'bottom':
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), sz)        # eighths of a point
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), color)
        else:
            el.set(qn('w:val'), 'none')
            el.set(qn('w:sz'), '0')
            el.set(qn('w:color'), 'auto')
        borders.append(el)
    tblPr.append(borders)


def _fixed_layout(table, widths_twips):
    """Force a fixed-width table layout with the given column widths (twips)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    # table width
    old_w = tblPr.find(qn('w:tblW'))
    if old_w is not None:
        tblPr.remove(old_w)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'dxa')
    tblW.set(qn('w:w'), str(sum(widths_twips)))
    tblPr.append(tblW)
    # fixed layout
    old_layout = tblPr.find(qn('w:tblLayout'))
    if old_layout is not None:
        tblPr.remove(old_layout)
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    # grid columns
    old_grid = tbl.find(qn('w:tblGrid'))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement('w:tblGrid')
    for w in widths_twips:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        grid.append(gc)
    tblPr.addnext(grid)


def _apply_size(paragraph, pt):
    """Force a font size (pt) on every run already in the paragraph."""
    for r in paragraph.runs:
        r.font.size = Pt(pt)


def _spacer(doc, after_pt=6):
    """Add a thin empty paragraph so consecutive tables stay separate."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after_pt)
    run = p.add_run('')
    run.font.size = Pt(4)
    return p


def _split_inline_choices(text):
    """
    Split run-on options out of a question stem when the model failed to put
    them on separate lines, e.g.:
        "If x = 2, what is y? (A) 8.4 (B) 9.6 (C) 11.6 (D) 12.4"
    Returns (stem, [(letter, choice), ...]) or (text, []) if no clean
    sequential A/B/C/D run of options is found.
    """
    if not text:
        return text, []
    # locate the first "A)" or "(A)" that begins the option run
    m = re.search(r'\(?\bA\)', text)
    if not m:
        return text, []
    stem = text[:m.start()].strip().rstrip(':').strip()
    rest = text[m.start():]
    # split on each "(X)" / "X)" marker, keeping the letter
    parts = re.split(r'\(?([A-D])\)\s*', rest)
    # parts -> ['', 'A', '8.4 ', 'B', '9.6 ', 'C', '11.6 ', 'D', '12.4']
    pairs = []
    tokens = parts[1:]
    for i in range(0, len(tokens) - 1, 2):
        letter = tokens[i].upper()
        value = tokens[i + 1].strip()
        pairs.append((letter, value))
    # require a clean sequential run starting at A (at least A and B)
    letters = [p[0] for p in pairs]
    expected = ['A', 'B', 'C', 'D'][:len(letters)]
    if len(pairs) >= 2 and letters == expected and all(v for _, v in pairs):
        return stem, pairs
    return text, []


def _parse_practice_markdown(md_text):
    """
    Parse the generated practice-set markdown into structured data.
    Returns dict {meta, questions, answers} or None if it can't be parsed
    into at least one question.
    """
    if not md_text:
        return None

    lines = md_text.split('\n')

    # --- Title / subject ------------------------------------------------
    subject = 'MATH'
    title = 'Practice Set'
    for ln in lines:
        s = ln.strip()
        m = re.match(r'^#{1,6}\s*(.*)', s)
        if m and 'Manual Set' in m.group(1):
            raw = m.group(1)
            if '📘' in raw or 'English' in raw:
                subject = 'ENGLISH'
            elif '📐' in raw or 'Math' in raw:
                subject = 'MATH'
            after = raw.split('Manual Set:', 1)[-1].strip()
            after = clean_text_for_export(after).strip(' :—-')
            if after:
                title = after
            break

    # --- Split questions vs answer key ---------------------------------
    split_idx = len(lines)
    for idx, ln in enumerate(lines):
        s = ln.strip().lstrip('#').strip()
        s_clean = re.sub(r'[^A-Za-z ]', '', s).strip().lower()
        if s_clean.startswith('answer key') or 'answer key' in s_clean:
            split_idx = idx
            break
        if s.strip() in ('--- PAGE BREAK ---', 'PAGE BREAK'):
            # answer key usually follows the page break
            split_idx = idx
            break

    q_lines = lines[:split_idx]
    a_lines = lines[split_idx:]

    # --- Parse questions -----------------------------------------------
    q_start = re.compile(
        r'^\**\s*\[?(Easy|Medium|Hard)\]?\s*\.?\s*(\d+)[.)]\s*\**\s*(.*)$',
        re.IGNORECASE,
    )
    opt_re = re.compile(r'^\**([A-D])\**[).]\s*(.*)$')

    questions = []
    cur = None
    pending_figure = None
    in_code = False
    code_lang = ''
    code_buf = []

    def close_question():
        nonlocal cur
        if cur is not None:
            cur['text'] = cur['text'].strip()
            # If the model kept the options inline (e.g. "... ? (A) 8 (B) 9 ...")
            # instead of on separate lines, split them out now.
            if not cur['choices']:
                stem, inline = _split_inline_choices(cur['text'])
                if inline:
                    cur['text'] = stem
                    cur['choices'] = inline
            questions.append(cur)
            cur = None

    for ln in q_lines:
        s = ln.strip()

        # code / figure blocks
        if s.startswith('```'):
            if not in_code:
                in_code = True
                code_lang = s[3:].strip().lower()
                code_buf = []
            else:
                in_code = False
                if code_lang in ('python-figure', 'python'):
                    img = execute_figure_code('\n'.join(code_buf))
                    if img:
                        # Per the prompt, a figure is placed IMMEDIATELY BEFORE
                        # the question that references it, so hand it to the next
                        # question that starts rather than the current one.
                        pending_figure = img
                code_buf = []
                code_lang = ''
            continue
        if in_code:
            code_buf.append(ln.rstrip())
            continue

        if not s:
            continue
        # skip stray markers / section headings inside question area
        if s in ('---', '--- PAGE BREAK ---') or s.startswith('─'):
            continue

        qm = q_start.match(s)
        if qm:
            close_question()
            cur = {
                'difficulty': qm.group(1).upper(),
                'num': int(qm.group(2)),
                'text': qm.group(3).strip(),
                'choices': [],
                'figure': pending_figure,
            }
            pending_figure = None
            continue

        if cur is None:
            # heading line before first question (e.g. the "### Manual Set" title)
            continue

        om = opt_re.match(s)
        if om:
            cur['choices'].append((om.group(1).upper(), om.group(2).strip()))
            continue

        # continuation of question text (only if no options captured yet)
        if not cur['choices']:
            cur['text'] = (cur['text'] + ' ' + s).strip()
        else:
            # trailing prose after the options -> append to last choice
            letter, txt = cur['choices'][-1]
            cur['choices'][-1] = (letter, (txt + ' ' + s).strip())

    close_question()

    if not questions:
        return None

    # --- Parse answer key ----------------------------------------------
    ans_re = re.compile(
        r'^\**\s*(\d+)[.)]\s*\**\s*\[?(Easy|Medium|Hard)?\]?\s*\**\s*'
        r'([A-D])\b\s*[-–—:.]*\s*(.*)$',
        re.IGNORECASE,
    )
    answers = {}
    cur_a = None
    for ln in a_lines:
        s = ln.strip()
        if not s:
            continue
        low = re.sub(r'[^a-z ]', '', s.lower()).strip()
        if low.startswith('answer key'):
            continue
        if s in ('---', '--- PAGE BREAK ---') or s.startswith('#'):
            continue
        am = ans_re.match(s)
        if am:
            num = int(am.group(1))
            cur_a = {
                'num': num,
                'difficulty': (am.group(2) or '').upper(),
                'answer': am.group(3).upper(),
                'explanation': am.group(4).strip(),
            }
            answers[num] = cur_a
        elif cur_a is not None:
            cur_a['explanation'] = (cur_a['explanation'] + '\n' + s).strip()

    # fill in difficulty from questions where the answer key omitted it
    q_by_num = {q['num']: q for q in questions}
    ordered_answers = []
    for num in sorted(answers):
        a = answers[num]
        if not a['difficulty'] and num in q_by_num:
            a['difficulty'] = q_by_num[num]['difficulty']
        ordered_answers.append(a)

    return {
        'subject': subject,
        'title': title,
        'questions': questions,
        'answers': ordered_answers,
    }


def _add_body_figure(cell, img_bytes, first_para):
    """Add a centered figure image into a question body cell."""
    p = first_para if first_para is not None else cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    try:
        p.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(4.0))
    except Exception:
        pass
    return p


def _make_card(doc, badge_text, header_runs, badge_key):
    """
    Create a 2-row 'card' table: row0 = [badge | header], row1 = merged body.
    Returns (table, body_cell). Rows are kept together across page breaks.
    header_runs: list of (text, bold, size_pt, color_hex_or_None).
    """
    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _fixed_layout(tbl, [_BADGE_COL, _BODY_COL])
    _table_borders(tbl, show=True)

    badge = _BADGE.get(badge_key, _BADGE['EASY'])

    # row 0 – badge cell
    badge_cell = tbl.cell(0, 0)
    _shade_cell(badge_cell, badge['fill'])
    bp = badge_cell.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = bp.add_run(badge_text)
    br.bold = True
    br.font.size = Pt(8)
    br.font.color.rgb = RGBColor.from_string(badge['text'])

    # row 0 – header cell
    header_cell = tbl.cell(0, 1)
    hp = header_cell.paragraphs[0]
    hp.paragraph_format.left_indent = Inches(0.08)
    for txt, bold, size_pt, color in header_runs:
        r = hp.add_run(txt)
        r.bold = bold
        r.font.size = Pt(size_pt)
        if color:
            r.font.color.rgb = RGBColor.from_string(color)

    # keep header row glued to the body row and unsplittable
    _keep_with_next(bp)
    _keep_with_next(hp)
    _row_cant_split(tbl.rows[0])

    # row 1 – merged body cell
    body_cell = tbl.cell(1, 0).merge(tbl.cell(1, 1))
    _row_cant_split(tbl.rows[1])

    return tbl, body_cell


def build_elite_practice_docx(md_text):
    """Build the Elite Prep styled .docx (mirrors sample-word-file.docx)."""
    data = _parse_practice_markdown(md_text)
    if not data:
        return None

    questions = data['questions']
    answers = data['answers']
    subject = data['subject']
    title = data['title']

    doc = Document()

    # base style
    normal = doc.styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    for section in doc.sections:
        section.top_margin = Inches(0.63)
        section.bottom_margin = Inches(0.63)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    # ---- difficulty counts -------------------------------------------
    counts = {'EASY': 0, 'MEDIUM': 0, 'HARD': 0}
    for q in questions:
        counts[q['difficulty']] = counts.get(q['difficulty'], 0) + 1

    # short chapter label (e.g. "Ch. 2")
    chap_short = subject.title()
    cm = re.search(r'Chapter\s*(\d+)', title, re.IGNORECASE)
    if cm:
        chap_short = f"Ch. {cm.group(1)}"

    # ---- 1. Header banner --------------------------------------------
    # Printer-friendly: no dark fill, navy text on white with a rule under it.
    head = doc.add_table(rows=1, cols=1)
    _fixed_layout(head, [_CARD_WIDTH])
    _table_bottom_rule(head, _ELITE_NAVY, '12')
    hcell = head.cell(0, 0)
    p = hcell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Elite Prep')
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string(_ELITE_NAVY)
    p2 = hcell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f'SAT {subject} PRACTICE MANUAL')
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor.from_string(_ELITE_ANSWER)
    p3 = hcell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(clean_text_for_export(title))
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor.from_string(_ELITE_MUTED)
    _spacer(doc, 4)

    # ---- 2. Stats bar -------------------------------------------------
    stats = doc.add_table(rows=1, cols=3)
    _fixed_layout(stats, [_CARD_WIDTH // 3] * 3)
    _table_borders(stats, show=False)
    stat_items = [
        (str(len(questions)), 'Total Questions'),
        (f"{counts['EASY']} / {counts['MEDIUM']} / {counts['HARD']}",
         'Easy / Medium / Hard'),
        (chap_short, subject.title()),
    ]
    for ci, (big, small) in enumerate(stat_items):
        c = stats.cell(0, ci)
        _shade_cell(c, _ELITE_LIGHT)
        bp = c.paragraphs[0]
        bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        br = bp.add_run(big)
        br.bold = True
        br.font.size = Pt(13)
        br.font.color.rgb = RGBColor.from_string(_ELITE_NAVY)
        sp = c.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sp.add_run(small)
        sr.font.size = Pt(8)
    _spacer(doc, 8)

    # ---- 3. Practice Questions section header ------------------------
    sec = doc.add_table(rows=1, cols=1)
    _fixed_layout(sec, [_CARD_WIDTH])
    _table_bottom_rule(sec, _ELITE_NAVY, '8')
    scell = sec.cell(0, 0)
    sp = scell.paragraphs[0]
    sp.paragraph_format.left_indent = Inches(0.08)
    sr = sp.add_run('Practice Questions')
    sr.bold = True
    sr.font.size = Pt(13)
    sr.font.color.rgb = RGBColor.from_string(_ELITE_NAVY)
    sp2 = scell.add_paragraph()
    sp2.paragraph_format.left_indent = Inches(0.08)
    sr2 = sp2.add_run('Solve each problem, then check your work against the Answer Key.')
    sr2.font.size = Pt(9)
    sr2.font.color.rgb = RGBColor.from_string(_ELITE_MUTED)
    _spacer(doc, 6)

    # ---- 4. Question cards -------------------------------------------
    for q in questions:
        diff = q['difficulty']
        _tbl, body = _make_card(
            doc,
            badge_text=diff,
            header_runs=[(f"Question {q['num']}", True, 10.5, _ELITE_NAVY)],
            badge_key=diff,
        )

        first = body.paragraphs[0]
        used_first = False

        if q['figure']:
            _add_body_figure(body, q['figure'], first)
            used_first = True

        # question text
        qp = first if not used_first else body.add_paragraph()
        used_first = True
        qp.paragraph_format.space_after = Pt(5)
        _add_formatted_text(qp, q['text'])
        _apply_size(qp, 10.5)

        # options
        for letter, choice in q['choices']:
            op = body.add_paragraph()
            op.paragraph_format.left_indent = Inches(0.18)
            op.paragraph_format.space_before = Pt(1)
            op.paragraph_format.space_after = Pt(1)
            lr = op.add_run(f"{letter})  ")
            lr.bold = True
            lr.font.color.rgb = RGBColor.from_string(_ELITE_OPTION)
            _add_formatted_text(op, choice)
            _apply_size(op, 10)

        _spacer(doc, 6)

    # ---- 5. Answer Key section header --------------------------------
    if answers:
        doc.add_page_break()
        ak = doc.add_table(rows=1, cols=1)
        _fixed_layout(ak, [_CARD_WIDTH])
        _table_bottom_rule(ak, _ELITE_NAVY, '8')
        akcell = ak.cell(0, 0)
        ap = akcell.paragraphs[0]
        ap.paragraph_format.left_indent = Inches(0.08)
        ar = ap.add_run('Answer Key & Explanations')
        ar.bold = True
        ar.font.size = Pt(13)
        ar.font.color.rgb = RGBColor.from_string(_ELITE_NAVY)
        ap2 = akcell.add_paragraph()
        ap2.paragraph_format.left_indent = Inches(0.08)
        ar2 = ap2.add_run('Step-by-step reasoning for every question.')
        ar2.font.size = Pt(9)
        ar2.font.color.rgb = RGBColor.from_string(_ELITE_MUTED)
        _spacer(doc, 6)

        # ---- 6. Explanation cards ------------------------------------
        for a in answers:
            diff = a['difficulty'] or 'EASY'
            _tbl, body = _make_card(
                doc,
                badge_text=diff,
                header_runs=[
                    (f"Question {a['num']}  •  Correct Answer: ", True, 10.5, _ELITE_NAVY),
                    (a['answer'], True, 10.5, _ELITE_ANSWER),
                ],
                badge_key=diff,
            )
            _shade_cell(body, _ELITE_LIGHT)

            steps = []
            for s in a['explanation'].split('\n'):
                # drop leading markdown bullet markers (*, -, +, •) the model
                # sometimes adds to each explanation step
                s = re.sub(r'^\s*[\*\-\+•]\s+', '', s)
                if s.strip():
                    steps.append(s)
            if not steps:
                steps = ['']
            for si, step in enumerate(steps):
                ep = body.paragraphs[0] if si == 0 else body.add_paragraph()
                ep.paragraph_format.space_after = Pt(3)
                _add_formatted_text(ep, step)
                _apply_size(ep, 10)

            _spacer(doc, 6)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# Helper: Convert Markdown to Word (.docx) document
def convert_markdown_to_docx(md_text):
    """Elite Prep styled export with a safe fallback to the linear renderer."""
    if not md_text:
        return None
    try:
        styled = build_elite_practice_docx(md_text)
        if styled:
            return styled
    except Exception as e:
        print(f"Elite docx build failed, using fallback: {e}")
    return _convert_markdown_to_docx_linear(md_text)


def _convert_markdown_to_docx_linear(md_text):
    if not md_text: return None

    doc = Document()
    
    # Set default font for Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1f, 0x29, 0x37)
    
    # Configure heading styles
    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Arial'
        heading_style.font.color.rgb = RGBColor(0x0C, 0x1E, 0x41)
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    code_block_lang = ""
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Handle code block start/end
        if stripped.startswith('```'):
            if not in_code_block:
                # Starting a code block
                code_block_lang = stripped[3:].strip().lower()
                in_code_block = True
                code_block_lines = []
                i += 1
                continue
            else:
                # Ending a code block
                in_code_block = False
                if code_block_lang in ('python-figure', 'python'):
                    # Execute matplotlib code and embed image
                    code_str = '\n'.join(code_block_lines)
                    img_bytes = execute_figure_code(code_str)
                    if img_bytes:
                        img_buf = io.BytesIO(img_bytes)
                        doc.add_picture(img_buf, width=Inches(4.5))
                        # Center the image
                        last_para = doc.paragraphs[-1]
                        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    # Regular code block - add as formatted text
                    code_text = '\n'.join(code_block_lines)
                    if code_text.strip():
                        para = doc.add_paragraph()
                        run = para.add_run(code_text)
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                code_block_lang = ""
                code_block_lines = []
                i += 1
                continue
        
        if in_code_block:
            code_block_lines.append(line.rstrip())
            i += 1
            continue
        
        # Strip blockquote markers: "> text" → "text", ">" → ""
        while stripped.startswith('>'):
            stripped = stripped[1:].strip()
        
        # Empty line → skip (no blank rows)
        if not stripped:
            i += 1
            continue
        
        # Convert LaTeX in the line
        stripped = _latex_to_readable(stripped)
        
        # Horizontal rule / Page Break
        if stripped in ('---', '--- PAGE BREAK ---') or stripped.startswith('─'):
            para = doc.add_paragraph()
            run = para.add_run('━' * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            i += 1
            continue
        
        # Headers (### > ## > #)
        if stripped.startswith('### '):
            heading = doc.add_heading(level=3)
            _add_formatted_text(heading, stripped[4:])
            i += 1
            continue
        elif stripped.startswith('## '):
            heading = doc.add_heading(level=2)
            _add_formatted_text(heading, stripped[3:])
            i += 1
            continue
        elif stripped.startswith('# '):
            heading = doc.add_heading(level=1)
            _add_formatted_text(heading, stripped[2:])
            i += 1
            continue
        
        # Bullet points: * or - at start
        if re.match(r'^[\*\-]\s+', stripped):
            content = re.sub(r'^[\*\-]\s+', '', stripped)
            para = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(para, content)
            i += 1
            continue
        
        # Numbered items: "1." or "1)" at start — render as text to preserve original numbering
        num_match = re.match(r'^(\d+)([.)\t])\s*(.*)', stripped)
        if num_match:
            num = num_match.group(1)
            content = num_match.group(3)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.2)
            run = para.add_run(f"{num}. ")
            run.bold = True
            _add_formatted_text(para, content)
            i += 1
            continue
        
        # Answer option lines: A) B) C) D) — with slight indent
        option_match = re.match(r'^([A-D])\)\s+(.*)', stripped)
        if option_match:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.3)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(f"{option_match.group(1)})  ")
            run.bold = True
            _add_formatted_text(para, option_match.group(2))
            i += 1
            continue
        
        # Regular paragraph
        para = doc.add_paragraph()
        _add_formatted_text(para, stripped)
        i += 1
    
    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def _add_formatted_text(paragraph, text):
    """Add text to a paragraph with bold and italic formatting support."""
    if not text:
        return
    
    # Remove Hangul/Emojis for Word consistency
    text = clean_text_for_export(text)
    
    # First convert any remaining LaTeX
    text = _latex_to_readable(text)
    
    # Split by bold markers **...**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            content = part[2:-2]
            # Handle superscripts inside bold
            _add_superscripts(paragraph, content, True)
        else:
            # Handle italic *...*
            italic_parts = re.split(r'(\*[^*]+\*)', part)
            for ip in italic_parts:
                if not ip:
                    continue
                if ip.startswith('*') and ip.endswith('*') and len(ip) > 2:
                    content = ip[1:-1]
                    _add_superscripts(paragraph, content, False, True)
                else:
                    _add_superscripts(paragraph, ip, False, False)

def _add_superscripts(paragraph, text, bold=False, italic=False):
    """Helper to add text with real superscript formatting for Word."""
    # Replace unicode superscripts with ^ marker for parsing
    text = text.replace('²', '^2').replace('³', '^3').replace('¹', '^1')
    text = text.replace('₀', '_0').replace('₁', '_1').replace('₂', '_2').replace('₃', '_3')
    
    # Split by ^ and _
    parts = re.split(r'(\^[0-9a-zA-Z]+|_[0-9a-zA-Z]+)', text)
    for part in parts:
        if part.startswith('^'):
            run = paragraph.add_run(part[1:])
            run.font.superscript = True
            run.bold = bold
            run.italic = italic
        elif part.startswith('_'):
            run = paragraph.add_run(part[1:])
            run.font.subscript = True
            run.bold = bold
            run.italic = italic
        else:
            run = paragraph.add_run(part)
            run.bold = bold
            run.italic = italic

class ElitePDF(FPDF):
    def header(self):
        # Modern Header with Navy bar
        self.set_fill_color(12, 30, 65) # Elite Navy
        self.rect(0, 0, 210, 15, 'F')
        self.set_font("Helvetica", "B", 10)
        self.set_text_color(255, 255, 255)
        self.set_y(5)
        self.cell(0, 5, "ELITE PREP | SAT ADAPTIVE PRACTICE STATION", align="C", ln=True)
        self.ln(10)

    def footer(self):
        # Subtle Footer
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(100, 100, 100)
        self.cell(0, 10, f"Confidential Elite Prep - Page {self.page_no()}/{{nb}}", align="C")

def convert_markdown_to_pdf(md_text):
    if not md_text: return None
    
    pdf = ElitePDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    
    elite_navy = (12, 30, 65)
    
    # Block grouping to keep questions together
    raw_lines = md_text.split('\n')
    blocks = []
    temp_block = []
    
    for line in raw_lines:
        s = line.strip()
        # Detect new sections or questions to start a new block
        # Match question numbers like 1. or **1.**
        is_q = re.match(r'^(\*\*|)(\d+)[.)]', s)
        if s.startswith('### ') or s.startswith('## ') or s.startswith('# ') or \
           is_q or s.startswith('Question ') or \
           s.startswith('Answer Key') or s == '--- PAGE BREAK ---':
            if temp_block: blocks.append(temp_block)
            temp_block = [line]
        else:
            temp_block.append(line)
    if temp_block: blocks.append(temp_block)
    
    for block in blocks:
        # Estimate height to prevent messy breaks
        est_h = len(block) * 8
        if any('```python-figure' in l for l in block): est_h += 80
        if pdf.get_y() + est_h > (pdf.h - 25): pdf.add_page()
        
        in_code = False
        code_lines = []
        code_lang = ""
        
        for line in block:
            pdf.set_x(18) # Relaxed margins
            s = line.strip()
            
            if s.startswith('```'):
                if not in_code:
                    in_code = True; code_lang = s[3:].strip().lower(); code_lines = []; continue
                else:
                    in_code = False
                    if code_lang in ('python-figure', 'python'):
                        img = execute_figure_code('\n'.join(code_lines))
                        if img: pdf.image(io.BytesIO(img), x=35, w=140); pdf.ln(5)
                    else:
                        pdf.set_font("Courier", size=9); pdf.set_fill_color(245, 245, 245)
                        pdf.multi_cell(0, 5, clean_text_for_export('\n'.join(code_lines)), fill=True)
                        pdf.set_font("Helvetica", size=11)
                    continue
            if in_code: code_lines.append(line); continue
                
            # Process text and clean markdown artifacts
            text = _latex_to_readable(s)
            while text.startswith('>'): text = text[1:].strip()
            
            # Regex for SAT elements
            q_m = re.match(r'^(\*\*|)(\d+)[.)\s\*]+(.*)', text)
            o_m = re.match(r'^(\*\*|)([A-D])\)[.\s\*]*(.*)', text)
            
            if text.startswith('### '):
                pdf.set_text_color(*elite_navy); pdf.set_font("Helvetica", "B", 14)
                pdf.cell(0, 12, clean_text_for_export(text[4:]), ln=True)
                pdf.set_font("Helvetica", size=11); pdf.set_text_color(0)
            elif text.startswith('## ') or text.startswith('# '):
                pdf.set_text_color(*elite_navy); pdf.set_font("Helvetica", "B", 18)
                label = text[text.find(' ')+1:]
                pdf.cell(0, 16, clean_text_for_export(label), ln=True)
                pdf.set_font("Helvetica", size=11); pdf.set_text_color(0)
            elif text in ('---', '--- PAGE BREAK ---'):
                if text == '--- PAGE BREAK ---': pdf.add_page()
                else:
                    pdf.ln(2); pdf.set_draw_color(200, 200, 200)
                    pdf.line(20, pdf.get_y(), 190, pdf.get_y()); pdf.ln(5)
            elif q_m:
                pdf.set_font("Helvetica", "B", 11); pdf.set_text_color(*elite_navy)
                num = q_m.group(2); content = q_m.group(3).strip().replace('**', '')
                pdf.write(7, f"{num}. "); pdf.set_font("Helvetica", size=11); pdf.set_text_color(0)
                pdf.multi_cell(0, 7, clean_text_for_export(content))
            elif o_m:
                pdf.set_x(28); pdf.set_font("Helvetica", "B", 10)
                letter = o_m.group(2); o_text = o_m.group(3).strip().replace('**', '')
                pdf.write(6, f"{letter}) "); pdf.set_font("Helvetica", size=11)
                pdf.multi_cell(0, 6, clean_text_for_export(o_text))
            else:
                cleaned = clean_text_for_export(text).replace('**', '').replace('__', '')
                if cleaned.strip(): pdf.multi_cell(0, 6, cleaned)
                else: pdf.ln(3)

                
    return bytes(pdf.output())

# Helper: Load Local Resources (JPGs)
def load_local_resources():
    resources = []
    content_parts = []
    current_dir = Path(".")
    
    # Core Textbooks
    file_map = {
        "SAT Math Textbook": "SAT-Math.jpg",
        "SAT English Textbook": "SAT-English.jpg",
        "Vocabulary List": ["Vocabulary.jpg", "Vocabulary.pdf"]
    }

    # DSAT Test Packets
    dsat_files = sorted(list(current_dir.glob("DSAT*.jpg")))
    
    # Load Core Files
    for label, filenames in file_map.items():
        if isinstance(filenames, str):
            filenames = [filenames]
        for fname in filenames:
            file_path = current_dir / fname
            if file_path.exists():
                try:
                    # Check session state/cache here normally, but simple upload for now
                    # (In prod, cache the File API URIs)
                    uploaded_file = upload_to_gemini(str(file_path))
                    if uploaded_file:
                        content_parts.append(uploaded_file)
                        resources.append(f"✅ {label} Loaded")
                        break
                except Exception:
                     resources.append(f"❌ Error loading {fname}")

    # Load Test Packets
    if dsat_files:
        resources.append(f"📚 {len(dsat_files)} Test Packets Detected")
        for dsat_path in dsat_files:
            try:
                uploaded_file = upload_to_gemini(str(dsat_path))
                if uploaded_file:
                    content_parts.append(uploaded_file)
            except Exception:
                pass

    return resources, content_parts, len(list(current_dir.glob("SAT*.jpg"))) # Return count for core


# --- Main Content ---
st.markdown("<h1 class='main-header'>SAT Practice Problem Generator</h1>", unsafe_allow_html=True)


# Initialization
if 'analysis_done' not in st.session_state:
    st.session_state.analysis_done = False
if 'full_report' not in st.session_state:
    st.session_state.full_report = ""
if 'weak_topics' not in st.session_state:
    st.session_state.weak_topics = None
if 'practice_result' not in st.session_state:
    st.session_state.practice_result = ""
if 'all_topics' not in st.session_state:
    st.session_state.all_topics = None
if 'manual_practice_result' not in st.session_state:
    st.session_state.manual_practice_result = ""

# Main Content
st.markdown("### 📚 Elite Topic Bank")
st.markdown("Select any Lesson Subject from the textbooks to generate new practice questions manually.")

# 1. Hardcoded Topic Lists (Optimized)
math_topics = [
    "Chapter 2: Solving Linear Equations", "Chapter 3: Systems of Linear Equations", "Chapter 4: Systems of Linear Equations in Context",
    "Chapter 5: Linear Functions", "Chapter 6: Graphs of Lines", "Chapter 7: Linear Models", "Chapter 8: Linear Inequalities",
    "Chapter 9: Absolute Value", "Chapter 10: Ratios", "Chapter 11: Units", "Chapter 12: Percentages", "Chapter 13: Percent Change",
    "Chapter 14: Rates", "Chapter 15: Lines and Angles", "Chapter 16: Similar Triangles", "Chapter 17: Measures of Center",
    "Chapter 18: Measures of Spread", "Chapter 19: Tables", "Chapter 20: Bar Graphs", "Chapter 21: Time-Series Graphs",
    "Chapter 22: Probability", "Chapter 23: Scatterplots", "Chapter 24: Inferences from Sample Statistics", "Chapter 25: Statistical Claims",
    "Chapter 26: Right Triangles", "Chapter 27: Trigonometry", "Chapter 28: Area and Volume", "Chapter 29: Circles",
    "Chapter 30: Completing the Square", "Chapter 31: Equations of Circles", "Chapter 32: Quadratic Equations", "Chapter 33: Rational Functions",
    "Chapter 34: Exponents and Radicals", "Chapter 35: Polynomials", "Chapter 36: Graphs of Functions", "Chapter 37: Nonlinear Equations in Context",
    "Chapter 38: Composite Functions", "Chapter 39: Linear and Exponential Models", "Chapter 40: Equivalent Expressions"
]

# Lesson subsections per Math chapter (from the textbook table of contents).
# Chapters that only contain a Problem Set (no numbered subsections) are
# omitted and generate at the whole-chapter level.
math_subtopics = {
    "Chapter 2: Solving Linear Equations": [
        "2.1 Plugging In Answers",
    ],
    "Chapter 3: Systems of Linear Equations": [
        "3.1 Finding the Solutions", "3.2 Solving for Unusual Quantities",
        "3.3 The Number of Solutions",
    ],
    "Chapter 6: Graphs of Lines": [
        "6.1 Equations", "6.2 Slope",
    ],
    "Chapter 7: Linear Models": [
        "7.1 Working without an equation", "7.2 Using a given equation",
    ],
    "Chapter 8: Linear Inequalities": [
        "8.1 Linear Inequality Basics", "8.2 Linear Inequalities in Word Problems",
        "8.3 Pairs of Inequalities", "8.4 Inequalities in the Plane",
    ],
    "Chapter 9: Absolute Value": [
        "9.1 Absolute Value is Nonnegative", "9.2 Graphs of Absolute Values",
        "9.3 Solving Absolute-value Equations", "9.4 Absolute Values as Inequalities",
    ],
    "Chapter 10: Ratios": [
        "10.1 Parts of a Whole", "10.2 Direct Proportions", "10.3 Scales",
    ],
    "Chapter 12: Percentages": [
        "12.1 Translating Percentage Word Problems", "12.2 Percentages in Context",
    ],
    "Chapter 13: Percent Change": [
        "13.1 Absolute vs. Percent Change", "13.2 Calculating Percent Change",
    ],
    "Chapter 15: Lines and Angles": [
        "15.1 Triangles", "15.2 Angles of Intersecting Lines",
    ],
    "Chapter 16: Similar Triangles": [
        "16.1 Establishing Similarity", "16.2 Similarity in Right Triangles",
    ],
    "Chapter 17: Measures of Center": [
        "17.1 Mean", "17.2 Median", "17.3 Mode", "17.4 When the Average is Given",
    ],
    "Chapter 18: Measures of Spread": [
        "18.1 Range", "18.2 Standard Deviation", "18.3 Why Spread Matters",
    ],
    "Chapter 21: Time-Series Graphs": [
        "21.1 Calculations with Data Points", "21.2 Interpreting Graphs in Context",
    ],
    "Chapter 23: Scatterplots": [
        "23.1 Best Fit",
    ],
    "Chapter 25: Statistical Claims": [
        "25.1 Population and Sample", "25.2 Flaws in Surveys", "25.3 The Right Conclusions",
    ],
    "Chapter 26: Right Triangles": [
        "26.1 Pythagorean Theorem", "26.2 Angles of Measures 30, 45, and 60 Degrees",
    ],
    "Chapter 27: Trigonometry": [
        "27.1 SohCahToa", "27.2 Complementary Angles", "27.3 The Unit Circle",
    ],
    "Chapter 28: Area and Volume": [
        "28.1 Area of a Circle", "28.2 Areas of Rectangles and Triangles",
        "28.3 Volume of a Sphere", "28.4 Volume of a Rectangular Solid",
        "28.5 Volumes of Cylinders", "28.6 Pyramids and Cones",
    ],
    "Chapter 29: Circles": [
        "29.1 Parts of a Circle", "29.2 Circumference and Arc Length of a Circle",
        "29.3 Area of a Sector of a Circle", "29.4 Tangents to Circles",
    ],
    "Chapter 32: Quadratic Equations": [
        "32.1 FOILing", "32.2 Two Important Factorizations",
        "32.3 Solving with Differences of Squares", "32.4 Simplifying Square Roots",
        "32.5 The Quadratic Formula", "32.6 A Shortcut", "32.7 Quadratics in the Plane",
    ],
    "Chapter 33: Rational Functions": [
        "33.1 Undefined Terms", "33.2 Solving Equations with Rational Expressions",
    ],
    "Chapter 34: Exponents and Radicals": [
        "34.1 Exponent Rules", "34.2 Solving Radical Equations",
    ],
    "Chapter 35: Polynomials": [
        "35.1 Definition of a Polynomial", "35.2 Addition and Subtraction of Polynomials",
        "35.3 Multiplication of Monomials",
        "35.4 Multiplication of Polynomials: The Distributive Property",
        "35.5 Roots", "35.6 Factors", "35.7 Remainders", "35.8 Factoring Cubics",
    ],
    "Chapter 36: Graphs of Functions": [
        "36.1 Values of Functions", "36.2 Zeros on Graphs of Polynomials", "36.3 Signs",
    ],
    "Chapter 37: Nonlinear Equations in Context": [
        "37.1 Quadratics", "37.2 Other Equations",
    ],
    "Chapter 39: Linear and Exponential Models": [
        "39.1 Distinguishing Linear from Exponential",
    ],
    "Chapter 40: Equivalent Expressions": [
        "40.1 Equivalent Quadratics", "40.2 Differences of Squares and Perfect Squares",
        "40.3 Polynomials: Like Terms", "40.4 Solving for a Particular Variable",
    ],
}

english_topics = [
    "Chapter 2: Central Ideas", "Chapter 3: Parts of Speech", "Chapter 4: Phrases", "Chapter 5: Active Reading",
    "Chapter 6: Clauses", "Chapter 7: Appositives", "Chapter 8: Command of Evidence (Textual)", "Chapter 9: Subject-Verb Agreement",
    "Chapter 10: Inferences", "Chapter 11: Verb Tense and Time Reference", "Chapter 12: Words in Context",
    "Chapter 13: Possessive Nouns and Possessive Determiners", "Chapter 14: Parentheticals", "Chapter 15: Modifier Placement",
    "Chapter 16: Text Structure and Purpose", "Chapter 17: Transitions", "Chapter 18: Informational Graphics",
    "Chapter 19: Rhetorical Synthesis", "Chapter 20: Cross-Text Connections", "Chapter 21: Punctuation"
]

# Lesson subsections per English chapter (from the textbook table of contents).
# When a chapter has entries here, the user can narrow generation to a single
# subsection instead of the whole chapter. "Problem Set" items are exercise
# sets, not lesson concepts, so they are intentionally excluded.
english_subtopics = {
    "Chapter 2: Central Ideas": [
        "2.1 Identifying the Main Idea", "2.2 Identifying Wrong Answers",
        "2.3 Question Prompts for Main Idea Questions", "2.4 Opening Information",
    ],
    "Chapter 3: Parts of Speech": [
        "3.1 Noun", "3.2 Verb", "3.3 Adjective", "3.4 Adverb", "3.5 Preposition",
        "3.6 Article", "3.7 Pronoun", "3.8 Conjunction", "3.9 Interjection",
    ],
    "Chapter 4: Phrases": [
        "4.1 Noun Phrases", "4.2 Adjective Phrases", "4.3 Prepositional Phrases",
        "4.4 Other Phrase Types",
    ],
    "Chapter 5: Active Reading": [
        "5.1 Becoming an Active Reader", "5.2 Contrast Words", "5.3 Absolutes",
    ],
    "Chapter 6: Clauses": [
        "6.1 Subordinating Conjunctions (Subordinators)", "6.2 Relative Clauses",
    ],
    "Chapter 7: Appositives": [
        "7.1 Trailing Appositives",
    ],
    "Chapter 8: Command of Evidence (Textual)": [
        "8.1 Science Texts", "8.2 Fiction and Poetry Texts",
    ],
    "Chapter 9: Subject-Verb Agreement": [
        "9.1 The Basic Rule", "9.2 And", "9.3 Or (and Nor)", "9.4 Expletives",
        "9.5 Collective Nouns", "9.6 Relative Pronouns", "9.7 Physics and Statistics",
        "9.8 Parentheticals and Pseudo-Conjunctions", "9.9 Either and Neither",
    ],
    "Chapter 10: Inferences": [
        "10.1 Identifying Inferences Questions", "10.2 Breaking Down a Text: Science",
        "10.3 Breaking Down a Text: Humanities",
        "10.4 Where to Find Questions on Inferences",
    ],
    "Chapter 11: Verb Tense and Time Reference": [
        "11.1 Tense", "11.2 Switching Tenses", "11.3 The Perfect",
    ],
    "Chapter 12: Words in Context": [
        "12.1 Why is Vocabulary Important?", "12.2 How to Study Vocabulary",
        "12.3 Types of Questions",
    ],
    "Chapter 13: Possessive Nouns and Possessive Determiners": [
        "13.1 Possessive Nouns", "13.2 Possessive Determiners",
    ],
    "Chapter 15: Modifier Placement": [
        "15.1 Dangling Modifiers", "15.2 Modifiers Generally",
    ],
    "Chapter 16: Text Structure and Purpose": [
        "16.1 For the Whole Text", "16.2 For Parts of the Text",
    ],
    "Chapter 17: Transitions": [
        "17.1 ACCES: Addition", "17.2 ACCES: Contrast", "17.3 ACCES: Causal",
        "17.4 ACCES: Example", "17.5 ACCES: Sequence (Time)",
    ],
    "Chapter 18: Informational Graphics": [
        "18.1 No Calculations", "18.2 Subjective Words", "18.3 Question Prompts",
        "18.4 Unjustifiable Conclusions", "18.5 Graphic Type: Tables",
        "18.6 Graphic Types: Line Graphs", "18.7 Graphic Types: Bar Graphs",
    ],
    "Chapter 19: Rhetorical Synthesis": [
        "19.1 The Questions",
    ],
    "Chapter 20: Cross-Text Connections": [
        "20.1 Question Prompts",
    ],
    "Chapter 21: Punctuation": [
        "21.1 The Period", "21.2 The Semicolon", "21.3 The Colon",
        "21.4 The Apostrophe", "21.5 The Comma", "21.6 Parentheses",
        "21.7 Hyphens and Dashes",
    ],
}

# Pre-load topics if not present
if not st.session_state.all_topics:
    st.session_state.all_topics = {
        "Math": math_topics,
        "English": english_topics
    }

# 2. Difficulty Selector
st.markdown("---")
st.markdown("#### Difficulty Distribution")
st.markdown("Set how many questions to generate per difficulty level.")

diff_col1, diff_col2, diff_col3, diff_col4 = st.columns([2, 1, 1, 1])
with diff_col1:
    st.markdown(" ")  # spacer label
with diff_col2:
    easy_count = st.number_input("🟢 Easy", min_value=0, max_value=20, value=3, step=1, key="easy_count")
with diff_col3:
    med_count  = st.number_input("🟡 Medium", min_value=0, max_value=20, value=4, step=1, key="med_count")
with diff_col4:
    hard_count = st.number_input("🔴 Hard", min_value=0, max_value=20, value=3, step=1, key="hard_count")

total_count = easy_count + med_count + hard_count
if total_count == 0:
    st.warning("Please set at least 1 question in any difficulty level.")
else:
    st.info(f"Total: **{total_count} questions** — 🟢 {easy_count} Easy  |  🟡 {med_count} Medium  |  🔴 {hard_count} Hard")

st.markdown("---")

# 3. Selection & Generation
WHOLE_CHAPTER = "-- Whole Chapter --"
col_m, col_e = st.columns(2)

with col_m:
    st.subheader("📐 Math Topics")
    math_list = st.session_state.all_topics.get("Math", math_topics) # Fallback to hardcoded
    selected_math = st.selectbox("Select Math Topic", ["-- Select --"] + math_list)

    # Optional lesson sub-topic (subsection) — only for chapters that have them.
    math_sub_options = math_subtopics.get(selected_math, [])
    selected_math_sub = WHOLE_CHAPTER
    if math_sub_options:
        selected_math_sub = st.selectbox(
            "Select Sub-topic (optional)",
            [WHOLE_CHAPTER] + math_sub_options,
            key="math_subtopic",
        )
    elif selected_math != "-- Select --":
        st.caption("No sub-topics listed for this chapter yet — questions will cover the whole chapter.")

    math_use_sub = bool(math_sub_options) and selected_math_sub != WHOLE_CHAPTER
    # ASCII hyphen (not em dash) so the title survives latin-1 export cleanly.
    math_topic = f"{selected_math} - {selected_math_sub}" if math_use_sub else selected_math

    if st.button("Generate Math Questions", disabled=(selected_math=="-- Select --" or total_count == 0)):
        with st.spinner(f"Generating {total_count} Math Questions for {math_topic}..."):
            res_status, context_parts, _ = load_local_resources()
            variation_id = random.randint(1000, 9999)

            # When a sub-topic is chosen, focus every question on that one
            # subsection; otherwise keep the broad whole-chapter coverage.
            if math_use_sub:
                math_topic_line = (
                    f"Create **{total_count} SAT Math Practice Questions** that focus "
                    f"specifically on the sub-topic **'{selected_math_sub}'**, which is a subsection "
                    f"of **'{selected_math}'**."
                )
                math_subtopic_focus = f"""
            **SUB-TOPIC FOCUS (CRITICAL):**
            - EVERY question MUST test the specific skill of the sub-topic '{selected_math_sub}'.
            - Do NOT drift into other subsections of '{selected_math}'. Keep the tested concept fixed on '{selected_math_sub}'.
            - Vary the numbers and scenarios, but the math skill being tested must stay '{selected_math_sub}'.
"""
                math_diversity_line = (
                    f"- All questions test the SAME concept ('{selected_math_sub}'), but each must use "
                    f"DIFFERENT numbers, coefficients, and problem scenarios."
                )
            else:
                math_topic_line = f"Create **{total_count} SAT Math Practice Questions** for the topic: **'{selected_math}'**."
                math_subtopic_focus = ""
                math_diversity_line = "- Each question MUST test a DIFFERENT sub-skill or concept within this topic."

            prompt = f"""
            {math_topic_line}

            **Variation Seed: {variation_id}** — Use this seed to ensure COMPLETELY UNIQUE questions.
            {math_subtopic_focus}
            **DIFFICULTY DISTRIBUTION (CRITICAL — MUST FOLLOW EXACTLY):**
            - Generate EXACTLY {easy_count} [Easy] questions, {med_count} [Medium] questions, and {hard_count} [Hard] questions.
            - Every single question MUST begin with its difficulty label in brackets: [Easy], [Medium], or [Hard].
            - Format MUST be: **[Easy] 1.** Question text... or **[Medium] 2.** Question text...
            - Do NOT skip or omit the difficulty label on any question.

            **DIVERSITY RULES (CRITICAL):**
            {math_diversity_line}
            - Use DIFFERENT numbers, coefficients, and constants than typical textbook examples.
            - Mix word problems, pure algebra, graph-based, and real-world application scenarios.
            - Do NOT repeat question patterns from any previous generation.

            **LANGUAGE RULES:**
            - Output MUST be 100% in English.
            - Do NOT use any Korean characters (Hangul).

            **INSTRUCTIONS:**
            - Mimic the exact difficulty and style of the "Elite Prep Textbooks".

            **LATEX FORMATTING RULES (CRITICAL — FOLLOW EXACTLY TO AVOID BROKEN RENDERING):**
            - Wrap EVERY mathematical expression in single dollar signs for inline math: $...$ (e.g. $P(A) = \\frac{{18}}{{30}}$, $r = \\sqrt{{121}} = 11$).
            - NEVER put plain English words or full sentences inside $...$. Dollar signs are ONLY for math symbols, variables, numbers-in-formulas, and equations.
            - Write explanatory sentences as normal text OUTSIDE any $...$. Example (CORRECT): Let $A$ be the event that the token is blue. So $P(A) = \\frac{{18}}{{30}}$.
            - Example (WRONG — do NOT do this): $Let A be the event that the token is blue. P(A) = \\frac{{18}}{{30}}$
            - Every opening $ MUST have a matching closing $ on the SAME line. Count them — they must be balanced (even number per line).
            - Do NOT use a stray/standalone $ character anywhere. Do NOT use asterisks (*) as separators between sentences.
            - Keep each inline math span short: wrap one expression at a time, not a whole paragraph.

            **FIGURE/GRAPH INSTRUCTIONS (CRITICAL - MUST INCLUDE):**
            - Include **2-3 questions** that require a graph, chart, or geometric figure.
            - For each figure, provide matplotlib Python code inside a ```python-figure block.
            - The code MUST use these exact variables: plt (matplotlib.pyplot), np (numpy), buf (io.BytesIO buffer).
            - End the code with: plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
            - Place the ```python-figure block IMMEDIATELY BEFORE the question that references it.
            - Reference the figure: "The figure above shows..." or "Based on the graph above..."
            - Graph types: coordinate planes with lines/curves, scatter plots, bar charts, geometric shapes (triangles, circles, rectangles with labeled dimensions), function graphs.
            - Make figures clean, labeled, with grid lines where appropriate.

            **CRITICAL STRUCTURE:**
            1. **Questions 1-{total_count}**: List the questions clearly.
               - EVERY question MUST start with its difficulty tag: **[Easy] 1.** / **[Medium] 2.** / **[Hard] 3.**
               - **Multiple Choice**: Put each option (A, B, C, D) on a NEW LINE.
            2. **--- PAGE BREAK ---**
            3. **Answer Key & Explanations**: You MUST provide this section at the very end.
               - Format: "1. [Easy] A - Explanation..."

            Output in clean Markdown.
            """
            res = get_gemini_response(prompt, context_parts, temperature=0.85)
            if res:
                st.session_state.manual_practice_result = f"### 📐 Manual Set: {math_topic}\n\n" + res

with col_e:
    st.subheader("📘 English Topics")
    eng_list = st.session_state.all_topics.get("English", english_topics) # Fallback to hardcoded
    selected_eng = st.selectbox("Select English Topic", ["-- Select --"] + eng_list)

    # Optional lesson sub-topic (subsection) — only for chapters that have them.
    sub_options = english_subtopics.get(selected_eng, [])
    selected_sub = WHOLE_CHAPTER
    if sub_options:
        selected_sub = st.selectbox(
            "Select Sub-topic (optional)",
            [WHOLE_CHAPTER] + sub_options,
            key="eng_subtopic",
        )
    elif selected_eng != "-- Select --":
        st.caption("No sub-topics listed for this chapter yet — questions will cover the whole chapter.")

    use_sub = bool(sub_options) and selected_sub != WHOLE_CHAPTER
    # ASCII hyphen (not em dash) so the title survives latin-1 export cleanly.
    eng_topic = f"{selected_eng} - {selected_sub}" if use_sub else selected_eng

    if st.button("Generate English Questions", disabled=(selected_eng=="-- Select --" or total_count == 0)):
            with st.spinner(f"Generating {total_count} English Questions for {eng_topic}..."):
                res_status, context_parts, _ = load_local_resources()
                variation_id = random.randint(1000, 9999)

                # When a sub-topic is chosen, focus every question on that one
                # subsection; otherwise keep the broad whole-chapter coverage.
                if use_sub:
                    topic_line = (
                        f"Create **{total_count} SAT English Practice Questions** that focus "
                        f"specifically on the sub-topic **'{selected_sub}'**, which is a subsection "
                        f"of **'{selected_eng}'**."
                    )
                    subtopic_focus = f"""
                **SUB-TOPIC FOCUS (CRITICAL):**
                - EVERY question MUST test the specific skill of the sub-topic '{selected_sub}'.
                - Do NOT drift into other subsections of '{selected_eng}'. Keep the tested concept fixed on '{selected_sub}'.
                - Vary the passage contexts and sentence content, but the grammar/skill point being tested must stay '{selected_sub}'.
"""
                    diversity_line = (
                        f"- All questions test the SAME concept ('{selected_sub}'), but each must use a "
                        f"DIFFERENT passage context, genre, and sentence content."
                    )
                else:
                    topic_line = f"Create **{total_count} SAT English Practice Questions** for the topic: **'{selected_eng}'**."
                    subtopic_focus = ""
                    diversity_line = "- Each question MUST test a DIFFERENT sub-skill or concept within this topic."

                prompt = f"""
                {topic_line}

                **Variation Seed: {variation_id}** — Use this seed to ensure COMPLETELY UNIQUE questions.
                {subtopic_focus}
                **DIFFICULTY DISTRIBUTION (CRITICAL — MUST FOLLOW EXACTLY):**
                - Generate EXACTLY {easy_count} [Easy] questions, {med_count} [Medium] questions, and {hard_count} [Hard] questions.
                - Every single question MUST begin with its difficulty label in brackets: [Easy], [Medium], or [Hard].
                - Format MUST be: **[Easy] 1.** Question text... or **[Medium] 2.** Question text...
                - Do NOT skip or omit the difficulty label on any question.

                **DIVERSITY RULES (CRITICAL):**
                {diversity_line}
                - Use DIFFERENT passage topics, genres, and writing styles (science, humanities, social science, literature).
                - Vary sentence complexity and vocabulary level across questions.
                - Do NOT repeat passage themes or question patterns from any previous generation.

                **LANGUAGE RULES:**
                - Output MUST be 100% in English.
                - Do NOT use any Korean characters (Hangul).

                **INSTRUCTIONS:**
                - Mimic the exact passage length and question style of the "Elite Prep Textbooks" / DSAT.

                **LATEX FORMATTING RULES (CRITICAL — FOLLOW EXACTLY TO AVOID BROKEN RENDERING):**
                - Wrap EVERY mathematical expression in single dollar signs for inline math: $...$ (e.g. $\\frac{{18}}{{30}}$).
                - NEVER put plain English words or full sentences inside $...$. Dollar signs are ONLY for math symbols, variables, numbers-in-formulas, and equations.
                - Every opening $ MUST have a matching closing $ on the SAME line (even number of $ per line).
                - Do NOT use a stray/standalone $ character. Do NOT use asterisks (*) as separators between sentences.

                **FIGURE/GRAPH INSTRUCTIONS:**
                - If the topic involves data interpretation or informational graphics, include 2-3 questions with figures.
                - For each figure, provide matplotlib Python code inside a ```python-figure block.
                - The code MUST use these exact variables: plt (matplotlib.pyplot), np (numpy), buf (io.BytesIO buffer).
                - End with: plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
                - Place the ```python-figure block IMMEDIATELY BEFORE the question that uses it.
                - Reference it: "Based on the figure above..." or "The graph above shows..."
                - Supported types: bar charts, line graphs, pie charts, scatter plots, data tables as charts.

                **CRITICAL STRUCTURE:**
                1. **Questions 1-{total_count}**: List the questions clearly.
                   - EVERY question MUST start with its difficulty tag: **[Easy] 1.** / **[Medium] 2.** / **[Hard] 3.**
                   - **Multiple Choice**: Put each option (A, B, C, D) on a NEW LINE.
                2. **--- PAGE BREAK ---**
                3. **Answer Key & Explanations**: You MUST provide this section at the very end.
                   - Format: "1. [Easy] A - Explanation..."

                Output in clean Markdown.
                """
                res = get_gemini_response(prompt, context_parts, temperature=0.85)
                if res:
                    st.session_state.manual_practice_result = f"### 📘 Manual Set: {eng_topic}\n\n" + res

# Display Result
if st.session_state.manual_practice_result:
    # Force Formatting
    st.session_state.manual_practice_result = ensure_formatting(st.session_state.manual_practice_result)
    
    st.markdown("---")
    
    # Header + Copy Button Layout
    r_col1, r_col2 = st.columns([5, 2])
    with r_col1:
        st.markdown("### 📝 Generated Practice Set")
    with r_col2:
        # Download Buttons
        docx_data = convert_markdown_to_docx(st.session_state.manual_practice_result)
        pdf_data = convert_markdown_to_pdf(st.session_state.manual_practice_result)
        
        if docx_data:
            st.download_button(
                label="📄 Download for Word",
                data=docx_data,
                file_name="Elite_Practice_Set.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    # Render with figure support + difficulty badges
    segments = parse_response_with_figures(st.session_state.manual_practice_result)
    for seg in segments:
        if seg["type"] == "text":
            content = sanitize_math_delimiters(seg["content"])
            content = add_difficulty_badges(content)
            st.markdown(content, unsafe_allow_html=True)
        elif seg["type"] == "figure":
            st.image(seg["image"], use_container_width=True)
    
    if st.button("Clear Manual Result"):
        st.session_state.manual_practice_result = ""
        st.rerun()

